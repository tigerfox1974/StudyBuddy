"""
StudyBuddy Utility Functions
Dosya hash, cache kontrol ve diger yardimci fonksiyonlar
"""

import hashlib
import json
import re
import io
import zipfile
import os
from typing import Tuple, Optional
import logging
from datetime import datetime
from models import Document, Result, UsageStats, User, UserUsageStats, Subscription, Payment, db
from config import Config
from flask import render_template, url_for

logger = logging.getLogger(__name__)

def get_file_hash(file_content):
    """
    Dosya iceriginin MD5 hash'ini dondurur
    Ayni icerik = Ayni hash = Cache hit
    
    Args:
        file_content: Binary dosya icerigi
        
    Returns:
        32 karakterlik hex hash
    """
    return hashlib.md5(file_content).hexdigest()


def check_cache(file_hash, user_level, user_type, user_id=None):
    """
    Cache'de bu dosya var mi kontrol eder
    
    Args:
        file_hash: Dosyanin MD5 hash'i
        user_level: Kullanici seviyesi (elementary, high_school, etc.)
        user_type: Kullanici tipi (student, teacher)
        user_id: Kullanici ID'si (opsiyonel). None ise sadece file_hash, user_level ve user_type ile filtreleme yapilir.
                 None degilse kullanici-spesifik cache kontrolu yapilir.
        
    Returns:
        Result object veya None
        
    Note:
        - user_id None ise: Genel cache kontrolu (kullanici-spesifik degil)
        - user_id belirtilirse: Kullanici-spesifik cache kontrolu
    """
    query = Document.query.filter_by(
        file_hash=file_hash,
        user_level=user_level,
        user_type=user_type
    )
    
    # user_id belirtilmişse, kullanıcı filtresini ekle
    if user_id is not None:
        query = query.filter_by(user_id=user_id)
    
    document = query.first()
    
    if document and document.results:
        # Last accessed guncelle
        from datetime import datetime
        document.last_accessed = datetime.utcnow()
        db.session.commit()
        
        return document.results[0]
    
    return None


def save_to_cache(file_hash, filename, file_type, file_size, user_level, user_type, 
                  results_data, ai_model, token_used, processing_time, user_id):
    """
    Sonuclari cache'e (veritabanina) kaydeder
    
    Args:
        file_hash: Dosya hash'i
        filename: Orijinal dosya adi
        file_type: Dosya tipi (pdf, docx, etc.)
        file_size: Dosya boyutu (byte)
        user_level: Kullanici seviyesi
        user_type: Kullanici tipi
        results_data: Uretilen icerikler (dict)
        ai_model: Kullanilan AI modeli
        token_used: Harcanan token
        processing_time: Islem suresi (saniye)
        user_id: Kullanici ID'si
        
    Returns:
        Result object
    """
    # Once ayni kombinasyon var mi kontrol et
    existing_document = Document.query.filter_by(
        file_hash=file_hash,
        user_level=user_level,
        user_type=user_type,
        user_id=user_id
    ).first()
    
    if existing_document:
        # Mevcut document varsa, eski result'u sil ve yenisini olustur
        # (Kullanici ayni dosyayi tekrar yuklerse guncel sonuc almali)
        for old_result in existing_document.results:
            db.session.delete(old_result)
        document = existing_document
        # Son erisim zamanini guncelle
        from datetime import datetime
        document.last_accessed = datetime.utcnow()
    else:
        # Yeni document kaydi olustur
        document = Document(
            file_hash=file_hash,
            original_filename=filename,
            file_type=file_type,
            file_size=file_size,
            user_level=user_level,
            user_type=user_type,
            user_id=user_id
        )
        db.session.add(document)
    
    db.session.flush()  # ID'yi al
    
    # Result kaydi olustur
    result = Result(
        document_id=document.id,
        summary=results_data['summary'],
        multiple_choice=json.dumps(results_data['multiple_choice'], ensure_ascii=False),
        short_answer=json.dumps(results_data['short_answer'], ensure_ascii=False),
        fill_blank=json.dumps(results_data['fill_blank'], ensure_ascii=False),
        true_false=json.dumps(results_data['true_false'], ensure_ascii=False),
        flashcards=json.dumps(results_data['flashcards'], ensure_ascii=False),
        ai_model=ai_model,
        token_used=token_used,
        processing_time=processing_time
    )
    db.session.add(result)
    db.session.flush()
    
    return result


def parse_cached_result(result):
    """
    Cache'den gelen Result nesnesini parse eder
    
    Args:
        result: Result database nesnesi
        
    Returns:
        dict: Frontend'e gonderilecek format
    """
    return {
        'summary': result.summary,
        'multiple_choice': json.loads(result.multiple_choice),
        'short_answer': json.loads(result.short_answer),
        'fill_blank': json.loads(result.fill_blank),
        'true_false': json.loads(result.true_false),
        'flashcards': json.loads(result.flashcards),
        'from_cache': True,
        'cached_date': result.created_at.strftime('%d %B %Y, %H:%M'),
        'token_used': result.token_used,
        'processing_time': result.processing_time
    }


def estimate_tokens(text):
    """
    Metin icin tahmini token sayisini hesaplar
    Yaklasik: 1 token = 4 karakter (Turkce icin)
    
    Args:
        text: Analiz edilecek metin
        
    Returns:
        int: Tahmini token sayisi
    """
    return len(text) // 4


def get_user_documents(user_id, limit=10, user_level=None, user_type=None):
    """
    Kullanicinin daha once isledig dosyalari getirir
    
    Args:
        user_id: Kullanici ID'si (zorunlu)
        limit: Maksimum sonuc sayisi (None ise limit yok, tumunu getirir)
        user_level: Filtre icin seviye (opsiyonel)
        user_type: Filtre icin tip (opsiyonel)
        
    Returns:
        List of Document objects
    """
    query = Document.query.filter_by(user_id=user_id)
    
    if user_level:
        query = query.filter_by(user_level=user_level)
    if user_type:
        query = query.filter_by(user_type=user_type)
    
    query = query.order_by(Document.last_accessed.desc())
    
    # limit None ise limit uygulama
    if limit is not None:
        query = query.limit(limit)
    
    return query.all()


def cleanup_old_cache(days=30):
    """
    Belirtilen gun oncesinden eski kayitlari siler
    
    Args:
        days: Kac gun oncesi (varsayilan 30)
        
    Returns:
        int: Silinen kayit sayisi
    """
    from datetime import datetime, timedelta
    
    cutoff_date = datetime.utcnow() - timedelta(days=days)
    old_documents = Document.query.filter(Document.last_accessed < cutoff_date).all()
    
    count = len(old_documents)
    for doc in old_documents:
        db.session.delete(doc)
    
    db.session.commit()
    return count


def validate_email_address(email):
    """
    Email adresini validate eder ve normalize eder
    
    Args:
        email: Email adresi string
        
    Returns:
        str: Normalize edilmiş email veya None (geçersizse)
    """
    if not email:
        return None
    
    try:
        from email_validator import validate_email, EmailNotValidError  # type: ignore
        # Test ortamında DNS kontrolü yapma (check_deliverability=False)
        # Production'da DNS kontrolü yapılabilir ama test ortamında sorun çıkarabilir
        validated = validate_email(email, check_deliverability=False)
        # validated.email deprecated, validated.normalized kullan
        return validated.normalized.lower().strip()
    except (EmailNotValidError, ImportError):
        return None


def generate_username_from_email(email):
    """
    Email'den otomatik username oluşturur
    
    Args:
        email: Email adresi
        
    Returns:
        str: Unique username
    """
    from models import User
    
    # Email'in @ öncesi kısmını al
    base_username = email.split('@')[0]
    
    # Özel karakterleri temizle (sadece alphanumeric ve underscore)
    base_username = re.sub(r'[^a-zA-Z0-9_]', '', base_username)
    
    # Uniqueness kontrolü
    username = base_username
    counter = 1
    while User.query.filter_by(username=username).first():
        username = f"{base_username}{counter}"
        counter += 1
    
    return username


def analyze_content_richness(text):
    """
    İçeriğin zenginliğini analiz eder ve yaklaşık soru sayısını tahmin eder
    
    Args:
        text: Döküman içeriği (string)
        
    Returns:
        dict: {
            'is_limited': bool,  # İçerik sınırlı mı? (<10 soru)
            'estimated_questions': int,  # Tahmini soru sayısı
            'word_count': int,  # Kelime sayısı
            'reason': str  # Sınırlı ise açıklama
        }
    """
    if not text or not text.strip():
        return {
            'is_limited': True,
            'estimated_questions': 0,
            'word_count': 0,
            'reason': 'Boş içerik'
        }
    
    # Kelime sayısını hesapla
    words = text.split()
    word_count = len(words)
    
    # Her 100 kelimeden yaklaşık 1 soru üretilebileceğini varsay
    # Bu değer AI modeline ve içeriğe göre değişebilir
    estimated_questions = max(1, word_count // 100)
    
    # 10'dan az soru üretilebilecekse içerik sınırlı sayılır
    is_limited = estimated_questions < 10
    
    reason = ''
    if is_limited:
        if word_count < 200:
            reason = 'Çok kısa içerik'
        elif word_count < 500:
            reason = 'Kısa içerik'
        else:
            reason = 'Orta uzunlukta içerik ama az soru üretilebilir'
    
    return {
        'is_limited': is_limited,
        'estimated_questions': estimated_questions,
        'word_count': word_count,
        'reason': reason
    }


def detect_main_topic(text):
    """
    Metinden ana konuyu/dersi tespit eder (basit keyword matching)
    Gelecekte daha gelişmiş AI tabanlı konu tespiti eklenebilir
    
    Args:
        text: Döküman içeriği (string)
        
    Returns:
        dict: {
            'subject': str,  # Tespit edilen ders (Matematik, Fizik vs.)
            'topic': str,  # Tespit edilen konu (Türev, Kuvvet vs.)
            'confidence': str  # 'high', 'medium', 'low'
        }
    """
    if not text:
        return {'subject': 'Genel', 'topic': 'Belirsiz', 'confidence': 'low'}
    
    text_lower = text.lower()
    
    # Basit keyword-based subject detection
    subject_keywords = {
        'Matematik': ['matematik', 'türev', 'integral', 'trigonometri', 'geometri', 'denklem', 'fonksiyon'],
        'Fizik': ['fizik', 'kuvvet', 'hız', 'ivme', 'enerji', 'momentum', 'elektrik', 'manyetizma'],
        'Kimya': ['kimya', 'atom', 'molekül', 'reaksiyon', 'element', 'periyodik', 'asit', 'baz'],
        'Biyoloji': ['biyoloji', 'hücre', 'doku', 'organ', 'canlı', 'ekosistem', 'genetik', 'evrim'],
        'Tarih': ['tarih', 'osmanlı', 'cumhuriyet', 'savaş', 'devlet', 'dönem', 'anlaşma'],
        'Coğrafya': ['coğrafya', 'iklim', 'nüfus', 'harita', 'kıta', 'deniz', 'ülke'],
        'Edebiyat': ['edebiyat', 'şiir', 'roman', 'öykü', 'nazım', 'nesir', 'edebi'],
        'İngilizce': ['english', 'grammar', 'vocabulary', 'tense', 'present', 'past', 'future']
    }
    
    detected_subject = 'Genel'
    max_match_count = 0
    
    for subject, keywords in subject_keywords.items():
        match_count = sum(1 for keyword in keywords if keyword in text_lower)
        if match_count > max_match_count:
            max_match_count = match_count
            detected_subject = subject
    
    confidence = 'high' if max_match_count >= 3 else ('medium' if max_match_count >= 1 else 'low')
    
    # Basit konu tespiti (ilk 200 karakter içinden)
    topic_preview = text[:200].strip()
    if len(topic_preview) > 50:
        topic = topic_preview[:50] + '...'
    else:
        topic = detected_subject + ' konusu'
    
    return {
        'subject': detected_subject,
        'topic': topic,
        'confidence': confidence
    }


def _read_head(content: bytes, n: int) -> bytes:
    """
    Dosya içeriğinin ilk n byte'ını güvenli şekilde okur
    
    Args:
        content: Dosya içeriği (bytes)
        n: Okunacak byte sayısı
        
    Returns:
        bytes: İlk n byte
    """
    return content[:n] if len(content) >= n else content


def _is_pdf(head: bytes) -> bool:
    """
    PDF dosyası olup olmadığını kontrol eder
    
    Args:
        head: Dosyanın ilk 5 byte'ı
        
    Returns:
        bool: PDF ise True
    """
    return head.startswith(b'%PDF-')


def _detect_office_zip(file_content: bytes) -> Optional[str]:
    """
    ZIP dosyasının DOCX mi PPTX mi olduğunu tespit eder
    
    Args:
        file_content: Dosya içeriği (bytes)
        
    Returns:
        str: 'docx', 'pptx' veya None
    """
    try:
        with zipfile.ZipFile(io.BytesIO(file_content), 'r') as zip_file:
            file_list = zip_file.namelist()
            
            # DOCX kontrolü: word/document.xml dosyası var mı?
            if any('word/document.xml' in name for name in file_list):
                return 'docx'
            
            # PPTX kontrolü: ppt/presentation.xml dosyası var mı?
            if any('ppt/presentation.xml' in name for name in file_list):
                return 'pptx'
            
            return None
    except (zipfile.BadZipFile, Exception):
        return None


def validate_file_signature(file_content: bytes, claimed_extension: str) -> Tuple[bool, Optional[str]]:
    """
    Dosya içeriğinin uzantısıyla eşleşip eşleşmediğini kontrol eder (Magic Number Validation)
    
    Args:
        file_content: Dosya içeriği (bytes)
        claimed_extension: Dosya uzantısı (örn: 'pdf', 'docx', 'pptx', 'txt')
        
    Returns:
        Tuple[bool, Optional[str]]: (True, None) geçerliyse, (False, error_message) geçersizse
    """
    if not file_content:
        return False, "Dosya boş veya geçersiz formatta."
    
    claimed_extension = claimed_extension.lower().lstrip('.')
    
    # TXT dosyaları için magic number kontrolü yok
    if claimed_extension == 'txt':
        return True, None
    
    # DOC (legacy) formatını reddet
    if claimed_extension == 'doc':
        return False, "Eski Word formatı (.doc) desteklenmiyor. Lütfen dosyanızı .docx formatına dönüştürün."
    
    # PDF kontrolü
    if claimed_extension == 'pdf':
        head = _read_head(file_content, 5)
        if not _is_pdf(head):
            return False, "Dosya içeriği uzantısıyla eşleşmiyor. PDF dosyası bekleniyor."
        return True, None
    
    # DOCX ve PPTX kontrolü (ZIP dosyaları)
    if claimed_extension in ('docx', 'pptx'):
        head = _read_head(file_content, 4)
        
        # ZIP signature kontrolü
        if not head.startswith(b'PK\x03\x04'):
            if claimed_extension == 'docx':
                return False, "Dosya içeriği uzantısıyla eşleşmiyor. Word belgesi bekleniyor."
            else:  # pptx
                return False, "Dosya içeriği uzantısıyla eşleşmiyor. PowerPoint sunumu bekleniyor."
        
        # ZIP içeriğini kontrol et
        detected_type = _detect_office_zip(file_content)
        
        if detected_type is None:
            return False, "Dosya bozuk veya geçersiz formatta."
        
        if detected_type != claimed_extension:
            if claimed_extension == 'docx':
                return False, "Dosya içeriği uzantısıyla eşleşmiyor. Word belgesi bekleniyor."
            else:  # pptx
                return False, "Dosya içeriği uzantısıyla eşleşmiyor. PowerPoint sunumu bekleniyor."
        
        return True, None
    
    # Desteklenmeyen uzantı
    return False, f"Desteklenmeyen dosya formatı: .{claimed_extension}"


def get_current_month_stats(user_id):
    """Kullanıcının bu ayki kullanım istatistiklerini getir veya oluştur"""
    now = datetime.utcnow()
    year = now.year
    month = now.month
    return UserUsageStats.get_or_create(user_id, year, month)


def increment_user_upload(user_id):
    """Kullanıcının dosya yükleme sayacını artır (cache miss durumunda)"""
    stats = get_current_month_stats(user_id)
    stats.increment_document()


def increment_user_cache_hit(user_id, tokens_saved):
    """Kullanıcının cache hit sayacını artır"""
    stats = get_current_month_stats(user_id)
    stats.increment_cache_hit(tokens_saved)


def check_user_upload_limit(user_id):
    """
    Kullanıcının aylık yükleme limitini kontrol et
    
    Returns:
        Tuple[bool, Optional[str], Optional[dict]]: 
        - (True, None, stats_dict): Limit OK, yükleme yapabilir
        - (False, error_message, stats_dict): Limit aşıldı, hata mesajı
    """
    user = User.query.get(user_id)
    if not user:
        return False, "Kullanıcı bulunamadı", None
    
    plan_type = user.subscription_plan or 'free'
    limit = Config.get_monthly_upload_limit(plan_type)
    
    # Premium plan (sınırsız)
    if limit is None:
        stats = get_current_month_stats(user_id)
        return True, None, {
            'plan': plan_type,
            'limit': None,
            'used': stats.get_total_uploads(),
            'remaining': None
        }
    
    # Limit kontrolü
    stats = get_current_month_stats(user_id)
    current_uploads = stats.get_total_uploads()
    
    if current_uploads >= limit:
        error_msg = f"Aylık dosya yükleme limitinize ulaştınız ({limit} dosya). Bir sonraki ayın 1'inde limit otomatik olarak sıfırlanacaktır."
        return False, error_msg, {
            'plan': plan_type,
            'limit': limit,
            'used': current_uploads,
            'remaining': 0
        }
    
    return True, None, {
        'plan': plan_type,
        'limit': limit,
        'used': current_uploads,
        'remaining': limit - current_uploads
    }


def get_user_stats_summary(user_id):
    """Kullanıcının genel kullanım özetini getir (dashboard için)"""
    user = User.query.get(user_id)
    if not user:
        return None
    
    stats = get_current_month_stats(user_id)
    can_upload, _, limit_info = check_user_upload_limit(user_id)
    
    total_documents = Document.query.filter_by(user_id=user_id).count()
    
    # Limit bilgisi için yüzde hesapla
    percentage = 0
    if limit_info and limit_info['limit']:
        percentage = (limit_info['used'] / limit_info['limit']) * 100
    
    # remaining değeri None olabilir (sınırsız planlar için)
    remaining = limit_info.get('remaining') if limit_info else None
    if remaining is None:
        # Sınırsız plan veya limit_info yok
        remaining = None
    else:
        remaining = int(remaining)
    
    return {
        'current_month': {
            'uploads': stats.documents_processed,
            'cache_hits': stats.cache_hits,
            'tokens_saved': stats.tokens_saved
        },
        'limit_info': {
            'plan': limit_info['plan'] if limit_info else 'free',
            'limit': limit_info['limit'] if limit_info else 5,
            'used': limit_info['used'] if limit_info else 0,
            'remaining': remaining,
            'percentage': percentage
        },
        'total_documents': total_documents
    }


def generate_invoice_pdf(payment, user, subscription_plan):
    """
    PDF fatura oluştur (ReportLab kullanarak)
    
    Args:
        payment: Payment object
        user: User object
        subscription_plan: Plan details dict
        
    Returns:
        str: Oluşturulan PDF dosyasının yolu
    """
    try:
        from reportlab.lib.pagesizes import A4  # type: ignore
        from reportlab.lib import colors  # type: ignore
        from reportlab.lib.units import cm  # type: ignore
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer  # type: ignore
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle  # type: ignore
        from reportlab.lib.enums import TA_CENTER, TA_RIGHT, TA_LEFT  # type: ignore
    except ImportError:
        raise ImportError("reportlab paketi yüklü değil. Lütfen 'pip install reportlab' komutu ile yükleyin.")
    
    # Invoice storage klasörünü oluştur
    invoice_dir = Config.INVOICE_STORAGE_PATH
    os.makedirs(invoice_dir, exist_ok=True)
    
    # PDF dosya yolu
    pdf_filename = f"{payment.invoice_number}.pdf"
    pdf_path = os.path.join(invoice_dir, pdf_filename)
    
    # PDF dokümanı oluştur
    doc = SimpleDocTemplate(pdf_path, pagesize=A4)
    story = []
    styles = getSampleStyleSheet()
    
    # Özel stiller
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#667eea'),
        spaceAfter=30,
        alignment=TA_CENTER
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=14,
        textColor=colors.HexColor('#333333'),
        spaceAfter=12
    )
    
    # Başlık
    story.append(Paragraph(Config.INVOICE_COMPANY_NAME, title_style))
    story.append(Spacer(1, 0.5*cm))
    
    # Fatura bilgileri
    invoice_data = [
        ['Fatura No:', payment.invoice_number],
        ['Tarih:', payment.created_at.strftime('%d %B %Y')],
        ['Durum:', 'Tamamlandı' if payment.status == 'completed' else payment.status]
    ]
    
    invoice_table = Table(invoice_data, colWidths=[4*cm, 10*cm])
    invoice_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTNAME', (1, 0), (1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
    ]))
    story.append(invoice_table)
    story.append(Spacer(1, 0.8*cm))
    
    # Müşteri bilgileri
    story.append(Paragraph('Müşteri Bilgileri', heading_style))
    customer_data = [
        ['Ad:', user.username],
        ['Email:', user.email]
    ]
    
    customer_table = Table(customer_data, colWidths=[4*cm, 10*cm])
    customer_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTNAME', (1, 0), (1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
    ]))
    story.append(customer_table)
    story.append(Spacer(1, 0.8*cm))
    
    # Ürün/Hizmet tablosu
    story.append(Paragraph('Fatura Detayları', heading_style))
    
    # Fiyat hesaplamaları
    subtotal = float(payment.amount)
    tax_amount = subtotal * (Config.INVOICE_TAX_RATE / 100)
    total = subtotal + tax_amount
    
    items_data = [
        ['Açıklama', 'Dönem', 'Tutar'],
        [
            subscription_plan.get('name', 'Premium Plan'),
            payment.billing_period,
            f"₺{subtotal:.2f}"
        ]
    ]
    
    items_table = Table(items_data, colWidths=[8*cm, 3*cm, 3*cm])
    items_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#667eea')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('ALIGN', (2, 0), (2, -1), 'RIGHT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('TOPPADDING', (0, 0), (-1, 0), 12),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 1), (-1, -1), 8),
        ('TOPPADDING', (0, 1), (-1, -1), 8),
        ('GRID', (0, 0), (-1, -1), 1, colors.grey),
    ]))
    story.append(items_table)
    story.append(Spacer(1, 0.5*cm))
    
    # Toplam tablosu
    if Config.INVOICE_TAX_RATE > 0:
        total_data = [
            ['Ara Toplam:', f"₺{subtotal:.2f}"],
            ['KDV (%{}):'.format(Config.INVOICE_TAX_RATE), f"₺{tax_amount:.2f}"],
            ['TOPLAM:', f"₺{total:.2f}"]
        ]
    else:
        total_data = [
            ['TOPLAM:', f"₺{subtotal:.2f}"]
        ]
    
    total_table = Table(total_data, colWidths=[11*cm, 3*cm])
    total_table.setStyle(TableStyle([
        ('ALIGN', (0, 0), (0, -1), 'RIGHT'),
        ('ALIGN', (1, 0), (1, -1), 'RIGHT'),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTNAME', (1, 0), (1, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 11),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('LINEABOVE', (0, -1), (-1, -1), 2, colors.HexColor('#667eea')),
    ]))
    story.append(total_table)
    story.append(Spacer(1, 1*cm))
    
    # Ödeme bilgileri
    if payment.payment_method:
        story.append(Paragraph(f'Ödeme Yöntemi: {payment.payment_method}', styles['Normal']))
    if payment.stripe_payment_intent_id:
        story.append(Paragraph(f'İşlem ID: {payment.stripe_payment_intent_id[:20]}...', styles['Normal']))
    
    story.append(Spacer(1, 1*cm))
    
    # Teşekkür mesajı
    thank_you = ParagraphStyle(
        'ThankYou',
        parent=styles['Normal'],
        fontSize=10,
        textColor=colors.grey,
        alignment=TA_CENTER,
        spaceBefore=20
    )
    story.append(Paragraph('Teşekkür ederiz!', thank_you))
    story.append(Paragraph(Config.INVOICE_COMPANY_ADDRESS, thank_you))
    
    # PDF'i oluştur
    doc.build(story)
    
    # Payment kaydına PDF yolunu kaydet (commit yapmaz, çağıran fonksiyon commit edecek)
    payment.invoice_pdf_path = pdf_path
    
    return pdf_path


def send_payment_confirmation_email(user_email, payment, invoice_pdf_path):
    """
    Ödeme onay email'i gönder (fatura PDF'i ile)
    
    Args:
        user_email: Email adresi
        payment: Payment object
        invoice_pdf_path: Invoice PDF dosya yolu
        
    Returns:
        bool: Başarılı ise True
    """
    try:
        from app import mail, app
        from flask_mail import Message  # type: ignore
        user = User.query.get(payment.user_id)
        plan = Config.SUBSCRIPTION_PLANS.get(payment.plan_type, {})
        dashboard_link = url_for('dashboard', _external=True)
        logo_url = url_for('static', filename='img/studybuddy-owl.png', _external=True)
        
        msg = Message(
            subject='StudyBuddy - Ödeme Onayı ve Fatura',
            recipients=[user_email],
            body=render_template('email/payment_success.txt', 
                               user=user, 
                               payment=payment, 
                               plan=plan,
                               dashboard_link=dashboard_link),
            html=render_template('email/payment_success.html', 
                               user=user, 
                               payment=payment, 
                               plan=plan,
                               dashboard_link=dashboard_link,
                               logo_url=logo_url)
        )
        
        # PDF'i ekle
        if os.path.exists(invoice_pdf_path):
            with app.open_resource(invoice_pdf_path) as pdf_file:
                msg.attach(
                    filename=f"{payment.invoice_number}.pdf",
                    content_type="application/pdf",
                    data=pdf_file.read()
                )
        
        mail.send(msg)
        return True
    except Exception as e:
        logger.error(f"Email gönderme hatası: {str(e)}")
        return False


def activate_user_subscription(user_id, plan_type, payment_id):
    """
    Kullanıcının premium aboneliğini aktif et
    
    Args:
        user_id: User ID
        plan_type: Plan tipi ('premium')
        payment_id: Payment ID
        
    Returns:
        Tuple[bool, Optional[Subscription], Optional[str]]: (success, subscription, error)
    """
    try:
        user = User.query.get(user_id)
        if not user:
            return False, None, "Kullanıcı bulunamadı"
        
        payment = Payment.query.get(payment_id)
        if not payment:
            return False, None, "Ödeme kaydı bulunamadı"
        
        # Idempotency: payment zaten işlenmiş ve subscription bağlıysa tekrar oluşturma
        if payment.status == 'completed' and payment.subscription_id:
            existing_sub = Subscription.query.filter_by(id=payment.subscription_id, user_id=user_id).first()
            if existing_sub:
                # Kullanıcının planını güncel kabul et
                user.subscription_plan = plan_type
                db.session.flush()
                return True, existing_sub, None
        
        # Kullanıcının planını güncelle
        user.subscription_plan = plan_type
        db.session.flush()
        
        # Yeni subscription kaydı oluştur
        subscription = Subscription(
            user_id=user_id,
            plan_type=plan_type,
            status='active',
            start_date=datetime.utcnow(),
            end_date=None  # Süresiz (monthly için)
        )
        db.session.add(subscription)
        db.session.flush()
        
        # Payment'ı subscription'a bağla
        payment.subscription_id = subscription.id
        payment.invoice_number = Payment.generate_invoice_number()
        db.session.flush()
        
        # Mevcut ayın usage stats'ını sıfırla (yeni fatura dönemi)
        # Veya mevcut stats'ı koru - burada koruyoruz
        # stats = get_current_month_stats(user_id)
        # stats.documents_processed = 0  # İsteğe bağlı: yeni dönemde sıfırla
        
        # Commit yapmıyoruz - çağıran fonksiyon commit edecek
        return True, subscription, None
        
    except Exception as e:
        # Rollback yapmıyoruz - çağıran fonksiyon rollback edecek
        return False, None, str(e)


def get_user_payment_history(user_id, limit=10):
    """
    Kullanıcının ödeme geçmişini getir
    
    Args:
        user_id: User ID
        limit: Maksimum kayıt sayısı
        
    Returns:
        List[Payment]: Payment nesneleri listesi (tarihe göre azalan)
    """
    return Payment.query.filter_by(user_id=user_id)\
        .order_by(Payment.created_at.desc())\
        .limit(limit)\
        .all()


def format_currency(amount, currency='TRY'):
    """
    Tutarı para birimi sembolü ile formatla
    
    Args:
        amount: Decimal veya float tutar
        currency: Para birimi kodu
        
    Returns:
        str: Formatlanmış tutar (örn: "₺49.99")
    """
    if currency == 'TRY':
        return f"₺{float(amount):.2f}"
    else:
        return f"{currency} {float(amount):.2f}"


# ============================================================================
# TOKEN (FİŞ) SİSTEMİ FONKSİYONLARI
# ============================================================================

def initialize_user_tokens(user):
    """
    Yeni kullanıcı için fiş sistemini başlat
    - 7 günlük deneme süresi başlat
    - Deneme fişlerini ver
    
    Args:
        user: User object
        
    Returns:
        None (user object güncellenir, commit yapılmaz)
    """
    from datetime import timedelta
    
    plan_config = Config.SUBSCRIPTION_PLANS.get('free', {})
    features = plan_config.get('features', {})
    
    # 7 günlük deneme başlat ve yalnızca bu anda deneme fişlerini ver
    if user.trial_ends_at is None:
        user.trial_ends_at = datetime.utcnow() + timedelta(days=features.get('trial_days', 7))
        user.tokens_remaining = features.get('trial_tokens', 10)


def is_trial_active(user):
    """
    Kullanıcının deneme süresi aktif mi kontrol et
    
    Args:
        user: User object
        
    Returns:
        bool: Deneme aktifse True
    """
    if not user.trial_ends_at:
        return False
    return datetime.utcnow() < user.trial_ends_at


def refresh_monthly_tokens(user):
    """
    Aylık fişleri yenile (abonelik planına göre)
    
    Args:
        user: User object
        
    Returns:
        None (user object güncellenir, commit yapılmaz)
    """
    from datetime import timedelta
    
    plan_config = Config.SUBSCRIPTION_PLANS.get(user.subscription_plan or 'free', {})
    features = plan_config.get('features', {})
    monthly_tokens = features.get('monthly_tokens', 3)
    
    # Son yenileme tarihini kontrol et
    now = datetime.utcnow()
    
    # İlk kez mi yoksa aylık yenileme zamanı mı?
    if user.last_token_refresh is None:
        # İlk kez - fişleri ver
        user.tokens_remaining = monthly_tokens
        user.last_token_refresh = now
    else:
        # Son yenilemeden 30 gün geçmiş mi?
        # Basit race mitigation: son 1 saat içinde yenilendiyse tekrar etme
        if (now - user.last_token_refresh).total_seconds() < 3600:
            return
        days_since_refresh = (now - user.last_token_refresh).days
        if days_since_refresh >= 30:
            # Aylık yenileme
            user.tokens_remaining = monthly_tokens
            user.last_token_refresh = now


def calculate_token_cost(question_types=None, include_export=False, user_plan='free'):
    """
    İşlem için gereken fiş maliyetini hesapla
    
    Args:
        question_types: List of question types to generate (None = all)
        include_export: Export yapılacak mı?
        user_plan: Kullanıcının planı
        
    Returns:
        float: Toplam fiş maliyeti
    """
    costs = Config.TOKEN_COSTS
    
    # Temel işleme (özet + flashcard)
    total = costs['base_processing']
    
    # Soru türleri için maliyet
    if question_types is None:
        # Tüm soru türleri
        question_types = ['multiple_choice', 'short_answer', 'fill_blank', 'true_false']
    
    total += len(question_types) * costs['question_type']
    
    # Export maliyeti
    if include_export:
        plan_config = Config.SUBSCRIPTION_PLANS.get(user_plan, {})
        features = plan_config.get('features', {})
        export_cost = features.get('export_cost_tokens', 2)
        total += export_cost
    
    return total


def check_user_tokens(user, required_tokens):
    """
    Kullanıcının yeterli fişi var mı kontrol et
    
    Args:
        user: User object
        required_tokens: Gereken fiş sayısı
        
    Returns:
        Tuple[bool, Optional[str], int]: (yeterli_mi, hata_mesajı, kalan_fiş)
    """
    # Önce aylık yenilemeyi kontrol et
    refresh_monthly_tokens(user)
    
    # Deneme süresi aktifse, deneme fişlerini kullan
    if is_trial_active(user):
        # Deneme süresinde - deneme fişleri kullanılabilir
        available_tokens = user.tokens_remaining
    else:
        # Deneme bitti - normal fişler
        available_tokens = user.tokens_remaining
    
    if available_tokens < required_tokens:
        error_msg = f"Yeterli fişiniz yok. Gereken: {required_tokens}, Mevcut: {available_tokens}. Lütfen fiş satın alın veya paket yükseltin."
        return False, error_msg, available_tokens
    
    return True, None, available_tokens


def deduct_tokens(user, amount):
    """
    Kullanıcıdan fiş düş
    
    Args:
        user: User object
        amount: Düşülecek fiş miktarı
        
    Returns:
        None (user object güncellenir, commit yapılmaz)
    """
    user.tokens_remaining = max(0, user.tokens_remaining - amount)


def add_tokens(user, amount):
    """
    Kullanıcıya fiş ekle (fiş satın alma için)
    
    Args:
        user: User object
        amount: Eklenecek fiş miktarı
        
    Returns:
        None (user object güncellenir, commit yapılmaz)
    """
    user.tokens_remaining += amount


def can_user_export(user):
    """
    Kullanıcı export yapabilir mi kontrol et
    
    Args:
        user: User object
        
    Returns:
        Tuple[bool, Optional[str], int]: (yapabilir_mi, hata_mesajı, gereken_fiş)
    """
    plan_config = Config.SUBSCRIPTION_PLANS.get(user.subscription_plan or 'free', {})
    features = plan_config.get('features', {})
    export_cost = features.get('export_cost_tokens', 2)
    
    # Premium plan'da export ücretsiz
    if export_cost == 0:
        return True, None, 0
    
    # Diğer planlarda fiş kontrolü
    can_afford, error_msg, available = check_user_tokens(user, export_cost)
    return can_afford, error_msg, export_cost


def get_user_token_info(user):
    """
    Kullanıcının fiş bilgilerini getir
    
    Args:
        user: User object
        
    Returns:
        dict: Fiş bilgileri
    """
    # Önce aylık yenilemeyi kontrol et
    refresh_monthly_tokens(user)
    
    plan_config = Config.SUBSCRIPTION_PLANS.get(user.subscription_plan or 'free', {})
    features = plan_config.get('features', {})
    monthly_tokens = features.get('monthly_tokens', 3)
    
    trial_active = is_trial_active(user)
    
    return {
        'tokens_remaining': user.tokens_remaining,
        'monthly_tokens': monthly_tokens,
        'trial_active': trial_active,
        'trial_ends_at': user.trial_ends_at.isoformat() if user.trial_ends_at else None,
        'last_refresh': user.last_token_refresh.isoformat() if user.last_token_refresh else None,
        'plan': user.subscription_plan or 'free'
    }


def generate_export_pdf(result, user, format_type='full'):
    """
    Result nesnesini PDF formatında sınav kağıdı olarak export et
    
    Args:
        result: Result model nesnesi (sorular, özet, flashcard'lar içerir)
        user: User model nesnesi (kullanıcı bilgileri için)
        format_type: 'full' (tüm içerik), 'questions_only' (sadece sorular), 'summary_only' (sadece özet)
        
    Returns:
        str: Oluşturulan PDF dosyasının yolu
    """
    try:
        from reportlab.lib.pagesizes import A4  # type: ignore
        from reportlab.lib import colors  # type: ignore
        from reportlab.lib.units import cm  # type: ignore
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak  # type: ignore
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle  # type: ignore
        from reportlab.lib.enums import TA_CENTER, TA_RIGHT, TA_LEFT  # type: ignore
    except ImportError:
        raise ImportError("reportlab paketi yüklü değil. Lütfen 'pip install reportlab' komutu ile yükleyin.")
    
    # Export storage klasörünü oluştur
    export_dir = Config.EXPORT_STORAGE_PATH
    os.makedirs(export_dir, exist_ok=True)
    
    # Dosya adı oluştur
    timestamp = datetime.utcnow().strftime('%Y%m%d_%H%M%S')
    base_filename = Config.EXPORT_FILENAME_FORMAT.format(result_id=result.id, timestamp=timestamp)
    pdf_filename = f"{base_filename}.pdf"
    pdf_path = os.path.join(export_dir, pdf_filename)
    
    # PDF dokümanı oluştur
    doc = SimpleDocTemplate(pdf_path, pagesize=A4)
    story = []
    styles = getSampleStyleSheet()
    
    # Özel stiller
    title_style = ParagraphStyle(
        'ExportTitle',
        parent=styles['Heading1'],
        fontSize=20,
        textColor=colors.HexColor('#2563EB'),
        spaceAfter=20,
        alignment=TA_CENTER
    )
    
    heading_style = ParagraphStyle(
        'ExportHeading',
        parent=styles['Heading2'],
        fontSize=14,
        textColor=colors.HexColor('#0F172A'),
        spaceAfter=12,
        spaceBefore=12
    )
    
    normal_style = ParagraphStyle(
        'ExportNormal',
        parent=styles['Normal'],
        fontSize=11,
        textColor=colors.HexColor('#0F172A'),
        spaceAfter=8
    )
    
    # Başlık
    story.append(Paragraph('StudyBuddy - Sınav Kağıdı', title_style))
    story.append(Spacer(1, 0.3*cm))
    
    # Dosya bilgileri
    info_data = [
        ['Dosya Adı:', result.document.original_filename],
        ['Oluşturulma Tarihi:', result.created_at.strftime('%d %B %Y, %H:%M')],
        ['Kullanıcı:', user.username]
    ]
    
    info_table = Table(info_data, colWidths=[4*cm, 10*cm])
    info_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTNAME', (1, 0), (1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
    ]))
    story.append(info_table)
    story.append(Spacer(1, 0.8*cm))
    
    # Özet bölümü (format_type'a göre)
    if format_type in ('full', 'summary_only'):
        story.append(Paragraph('Doküman Özeti', heading_style))
        summary_text = result.summary or 'Özet bulunamadı.'
        # Markdown formatını temizle (basit)
        summary_clean = summary_text.replace('**', '').replace('*', '').replace('#', '')
        story.append(Paragraph(summary_clean, normal_style))
        story.append(Spacer(1, 0.5*cm))
        story.append(PageBreak())
    
    if format_type == 'summary_only':
        doc.build(story)
        return pdf_path
    
    # Soru türlerini parse et
    try:
        multiple_choice = json.loads(result.multiple_choice) if result.multiple_choice else []
        short_answer = json.loads(result.short_answer) if result.short_answer else []
        fill_blank = json.loads(result.fill_blank) if result.fill_blank else []
        true_false = json.loads(result.true_false) if result.true_false else []
        flashcards = json.loads(result.flashcards) if result.flashcards else []
    except (json.JSONDecodeError, TypeError):
        multiple_choice = []
        short_answer = []
        fill_blank = []
        true_false = []
        flashcards = []
    
    # Çoktan Seçmeli Sorular
    if multiple_choice and format_type in ('full', 'questions_only'):
        story.append(Paragraph('Çoktan Seçmeli Sorular', heading_style))
        for idx, q in enumerate(multiple_choice, 1):
            question_text = q.get('question', 'Soru bulunamadı.')
            story.append(Paragraph(f"{idx}. {question_text}", normal_style))
            
            options = q.get('options', [])
            for opt_idx, option in enumerate(options, 1):
                option_label = chr(64 + opt_idx)  # A, B, C, D
                story.append(Paragraph(f"   {option_label}) {option}", normal_style))
            
            story.append(Spacer(1, 0.3*cm))
        story.append(Spacer(1, 0.5*cm))
    
    # Kısa Cevap Soruları
    if short_answer and format_type in ('full', 'questions_only'):
        story.append(Paragraph('Kısa Cevap Soruları', heading_style))
        for idx, q in enumerate(short_answer, 1):
            question_text = q.get('question', 'Soru bulunamadı.')
            story.append(Paragraph(f"{idx}. {question_text}", normal_style))
            story.append(Spacer(1, 0.3*cm))
        story.append(Spacer(1, 0.5*cm))
    
    # Boş Doldurma Soruları
    if fill_blank and format_type in ('full', 'questions_only'):
        story.append(Paragraph('Boş Doldurma Soruları', heading_style))
        for idx, q in enumerate(fill_blank, 1):
            question_text = q.get('question', 'Soru bulunamadı.')
            story.append(Paragraph(f"{idx}. {question_text}", normal_style))
            story.append(Spacer(1, 0.3*cm))
        story.append(Spacer(1, 0.5*cm))
    
    # Doğru-Yanlış Soruları
    if true_false and format_type in ('full', 'questions_only'):
        story.append(Paragraph('Doğru-Yanlış Soruları', heading_style))
        for idx, q in enumerate(true_false, 1):
            question_text = q.get('question', 'Soru bulunamadı.')
            story.append(Paragraph(f"{idx}. {question_text}", normal_style))
            story.append(Paragraph("   ( ) Doğru   ( ) Yanlış", normal_style))
            story.append(Spacer(1, 0.3*cm))
        story.append(Spacer(1, 0.5*cm))
    
    # Flashcard'lar
    if flashcards and format_type == 'full':
        story.append(PageBreak())
        story.append(Paragraph('Flashcard Seti', heading_style))
        
        # Flashcard'ları tablo halinde göster
        flashcard_data = [['Ön Yüz', 'Arka Yüz']]
        for card in flashcards:
            front = card.get('front', 'Ön yüz bulunamadı.')
            back = card.get('back', 'Arka yüz bulunamadı.')
            flashcard_data.append([front, back])
        
        flashcard_table = Table(flashcard_data, colWidths=[9*cm, 9*cm])
        flashcard_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2563EB')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 11),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
            ('TOPPADDING', (0, 0), (-1, 0), 10),
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 1), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 1), (-1, -1), 8),
            ('TOPPADDING', (0, 1), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey),
        ]))
        story.append(flashcard_table)
        story.append(Spacer(1, 0.5*cm))
    
    # Cevap Anahtarı (opsiyonel)
    if Config.EXPORT_INCLUDE_ANSWER_KEY and format_type in ('full', 'questions_only'):
        story.append(PageBreak())
        story.append(Paragraph('Cevap Anahtarı', heading_style))
        
        answer_key_data = []
        
        # Çoktan Seçmeli cevapları
        if multiple_choice:
            answer_key_data.append(['Çoktan Seçmeli Sorular', ''])
            for idx, q in enumerate(multiple_choice, 1):
                correct_answer = q.get('correct_answer', 'N/A')
                answer_key_data.append([f"{idx}.", correct_answer])
            answer_key_data.append(['', ''])  # Boş satır
        
        # Kısa Cevap cevapları
        if short_answer:
            answer_key_data.append(['Kısa Cevap Soruları', ''])
            for idx, q in enumerate(short_answer, 1):
                answer = q.get('answer', 'Cevap bulunamadı.')
                answer_key_data.append([f"{idx}.", answer])
            answer_key_data.append(['', ''])  # Boş satır
        
        # Boş Doldurma cevapları
        if fill_blank:
            answer_key_data.append(['Boş Doldurma Soruları', ''])
            for idx, q in enumerate(fill_blank, 1):
                answer = q.get('answer', 'Cevap bulunamadı.')
                answer_key_data.append([f"{idx}.", answer])
            answer_key_data.append(['', ''])  # Boş satır
        
        # Doğru-Yanlış cevapları
        if true_false:
            answer_key_data.append(['Doğru-Yanlış Soruları', ''])
            for idx, q in enumerate(true_false, 1):
                answer = 'Doğru' if q.get('answer', False) else 'Yanlış'
                answer_key_data.append([f"{idx}.", answer])
        
        if answer_key_data:
            answer_table = Table(answer_key_data, colWidths=[4*cm, 14*cm])
            answer_table.setStyle(TableStyle([
                ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                ('FONTNAME', (1, 0), (1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                ('TOPPADDING', (0, 0), (-1, -1), 6),
                ('GRID', (0, 0), (-1, -1), 1, colors.lightgrey),
            ]))
            story.append(answer_table)
    
    # PDF'i oluştur
    doc.build(story)
    
    return pdf_path


def generate_export_docx(result, user, format_type='full'):
    """
    Result nesnesini DOCX formatında sınav kağıdı olarak export et
    
    Args:
        result: Result model nesnesi (sorular, özet, flashcard'lar içerir)
        user: User model nesnesi (kullanıcı bilgileri için)
        format_type: 'full' (tüm içerik), 'questions_only' (sadece sorular), 'summary_only' (sadece özet)
        
    Returns:
        str: Oluşturulan DOCX dosyasının yolu
    """
    try:
        from docx import Document  # type: ignore
        from docx.shared import Pt, RGBColor, Inches  # type: ignore
        from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
    except ImportError:
        raise ImportError("python-docx paketi yüklü değil. Lütfen 'pip install python-docx' komutu ile yükleyin.")
    
    # Export storage klasörünü oluştur
    export_dir = Config.EXPORT_STORAGE_PATH
    os.makedirs(export_dir, exist_ok=True)
    
    # Dosya adı oluştur
    timestamp = datetime.utcnow().strftime('%Y%m%d_%H%M%S')
    base_filename = Config.EXPORT_FILENAME_FORMAT.format(result_id=result.id, timestamp=timestamp)
    docx_filename = f"{base_filename}.docx"
    docx_path = os.path.join(export_dir, docx_filename)
    
    # DOCX dokümanı oluştur
    doc = Document()
    
    # Başlık
    title = doc.add_heading('StudyBuddy - Sınav Kağıdı', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Dosya bilgileri
    info_para = doc.add_paragraph()
    info_para.add_run('Dosya Adı: ').bold = True
    info_para.add_run(result.document.original_filename)
    info_para.add_run('\nOluşturulma Tarihi: ').bold = True
    info_para.add_run(result.created_at.strftime('%d %B %Y, %H:%M'))
    info_para.add_run('\nKullanıcı: ').bold = True
    info_para.add_run(user.username)
    
    doc.add_paragraph()  # Boş satır
    
    # Özet bölümü (format_type'a göre)
    if format_type in ('full', 'summary_only'):
        doc.add_heading('Doküman Özeti', 1)
        summary_text = result.summary or 'Özet bulunamadı.'
        # Markdown formatını temizle (basit)
        summary_clean = summary_text.replace('**', '').replace('*', '').replace('#', '')
        doc.add_paragraph(summary_clean)
        doc.add_page_break()
    
    if format_type == 'summary_only':
        doc.save(docx_path)
        return docx_path
    
    # Soru türlerini parse et
    try:
        multiple_choice = json.loads(result.multiple_choice) if result.multiple_choice else []
        short_answer = json.loads(result.short_answer) if result.short_answer else []
        fill_blank = json.loads(result.fill_blank) if result.fill_blank else []
        true_false = json.loads(result.true_false) if result.true_false else []
        flashcards = json.loads(result.flashcards) if result.flashcards else []
    except (json.JSONDecodeError, TypeError):
        multiple_choice = []
        short_answer = []
        fill_blank = []
        true_false = []
        flashcards = []
    
    # Çoktan Seçmeli Sorular
    if multiple_choice and format_type in ('full', 'questions_only'):
        doc.add_heading('Çoktan Seçmeli Sorular', 1)
        for idx, q in enumerate(multiple_choice, 1):
            question_text = q.get('question', 'Soru bulunamadı.')
            para = doc.add_paragraph(f"{idx}. {question_text}", style='List Number')
            
            options = q.get('options', [])
            for opt_idx, option in enumerate(options, 1):
                option_label = chr(64 + opt_idx)  # A, B, C, D
                doc.add_paragraph(f"   {option_label}) {option}", style='List Bullet')
        doc.add_paragraph()  # Boş satır
    
    # Kısa Cevap Soruları
    if short_answer and format_type in ('full', 'questions_only'):
        doc.add_heading('Kısa Cevap Soruları', 1)
        for idx, q in enumerate(short_answer, 1):
            question_text = q.get('question', 'Soru bulunamadı.')
            doc.add_paragraph(f"{idx}. {question_text}", style='List Number')
        doc.add_paragraph()  # Boş satır
    
    # Boş Doldurma Soruları
    if fill_blank and format_type in ('full', 'questions_only'):
        doc.add_heading('Boş Doldurma Soruları', 1)
        for idx, q in enumerate(fill_blank, 1):
            question_text = q.get('question', 'Soru bulunamadı.')
            doc.add_paragraph(f"{idx}. {question_text}", style='List Number')
        doc.add_paragraph()  # Boş satır
    
    # Doğru-Yanlış Soruları
    if true_false and format_type in ('full', 'questions_only'):
        doc.add_heading('Doğru-Yanlış Soruları', 1)
        for idx, q in enumerate(true_false, 1):
            question_text = q.get('question', 'Soru bulunamadı.')
            para = doc.add_paragraph(f"{idx}. {question_text}", style='List Number')
            doc.add_paragraph("   ( ) Doğru   ( ) Yanlış")
        doc.add_paragraph()  # Boş satır
    
    # Flashcard'lar
    if flashcards and format_type == 'full':
        doc.add_page_break()
        doc.add_heading('Flashcard Seti', 1)
        
        # Flashcard'ları tablo halinde göster
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # Header
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Ön Yüz'
        header_cells[1].text = 'Arka Yüz'
        header_cells[0].paragraphs[0].runs[0].bold = True
        header_cells[1].paragraphs[0].runs[0].bold = True
        
        # Data rows
        for card in flashcards:
            row_cells = table.add_row().cells
            row_cells[0].text = card.get('front', 'Ön yüz bulunamadı.')
            row_cells[1].text = card.get('back', 'Arka yüz bulunamadı.')
    
    # Cevap Anahtarı (opsiyonel)
    if Config.EXPORT_INCLUDE_ANSWER_KEY and format_type in ('full', 'questions_only'):
        doc.add_page_break()
        doc.add_heading('Cevap Anahtarı', 1)
        
        # Çoktan Seçmeli cevapları
        if multiple_choice:
            doc.add_heading('Çoktan Seçmeli Sorular', 2)
            for idx, q in enumerate(multiple_choice, 1):
                correct_answer = q.get('correct_answer', 'N/A')
                doc.add_paragraph(f"{idx}. {correct_answer}", style='List Number')
            doc.add_paragraph()  # Boş satır
        
        # Kısa Cevap cevapları
        if short_answer:
            doc.add_heading('Kısa Cevap Soruları', 2)
            for idx, q in enumerate(short_answer, 1):
                answer = q.get('answer', 'Cevap bulunamadı.')
                doc.add_paragraph(f"{idx}. {answer}", style='List Number')
            doc.add_paragraph()  # Boş satır
        
        # Boş Doldurma cevapları
        if fill_blank:
            doc.add_heading('Boş Doldurma Soruları', 2)
            for idx, q in enumerate(fill_blank, 1):
                answer = q.get('answer', 'Cevap bulunamadı.')
                doc.add_paragraph(f"{idx}. {answer}", style='List Number')
            doc.add_paragraph()  # Boş satır
        
        # Doğru-Yanlış cevapları
        if true_false:
            doc.add_heading('Doğru-Yanlış Soruları', 2)
            for idx, q in enumerate(true_false, 1):
                answer = 'Doğru' if q.get('answer', False) else 'Yanlış'
                doc.add_paragraph(f"{idx}. {answer}", style='List Number')
    
    # DOCX'i kaydet
    doc.save(docx_path)
    
    return docx_path


